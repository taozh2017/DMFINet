import subprocess
from docx import Document

# 创建一个新的 Word 文档
doc = Document()

# 定义固定的命令参数
base_command = [
    'python', 'EN2CTTest.py',
    '--model', 'cyclegan_2D',
    '--which_model_netG', 'mymodel',
    '--which_direction', 'AtoB',
    '--dataset_mode', 'unaligned_2D',
    '--norm', 'instance',
    '--fineSize', '512',
    '--batchSize', '1',
    '--patch_size', '8',
    '--mask_ratio', '0.0',
    '--base_filter', '64',
    '--res_net'
]

# 定义一个函数来运行命令并捕获输出
def run_command(command, description):
    try:
        result = subprocess.run(command, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        output = result.stdout + result.stderr
        return output
    except subprocess.CalledProcessError as e:
        return str(e)

# 定义 which_epoch 的值列表
epochs = [80, 85, 90, 95, 100]

# 循环执行命令
for epoch in epochs:
    command1 = base_command + ['--which_epoch', str(epoch)]
    run_command(command1, f"Command 1: EN2CTTest with epoch {epoch}")

    # 运行第二个命令并捕获输出
    command2 = [
        'python', '-m', 'pytorch_fid',
        'result/cyclegan_2D/gt/', 'result/cyclegan_2D/output/'
    ]
    fid_output = run_command(command2, f"Command 2: Pytorch FID for epoch {epoch}")
    doc.add_heading(f"Pytorch FID for epoch {epoch}", level=1)
    doc.add_paragraph(fid_output)

# 保存 Word 文档
doc.save("output.docx")

print("Outputs have been saved to output.docx")
